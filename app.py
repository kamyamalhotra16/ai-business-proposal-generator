
import os
import pandas as pd

from docx import Document

from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser

from langchain_core.runnables import RunnableSequence
from langchain_core.runnables import RunnableParallel

from pydantic import BaseModel

print("All libraries imported successfully!")

GOOGLE_API_KEY="AQ.Ab8RN6InDk5b8KPDItVfcRnGf2dyh_0m3VBSR5wBvXlKSHgrTQ"

os.environ["GOOGLE_API_KEY"]=GOOGLE_API_KEY

from langchain_google_genai import ChatGoogleGenerativeAI

llm = ChatGoogleGenerativeAI(
    model="gemini-3.5-flash-lite",
    temperature=0.3
)

proposal_template = """

You are a professional business proposal writer.

Create a complete proposal using the information below.

Company Profile:

{company_profile}


Client Requirement:

{client_requirement}


The proposal must contain these sections:

1. Executive Summary

2. Company Introduction

3. Understanding of Client Requirements

4. Proposed Solution

5. Scope of Work

6. Deliverables

7. Implementation Timeline

8. Technology Stack

9. Pricing

10. Assumptions

11. Terms and Conditions

12. Conclusion


Write in professional business language.

"""


prompt = PromptTemplate(
    template=proposal_template,
    input_variables=[
        "company_profile",
        "client_requirement"
    ]
)

parser = StrOutputParser()


proposal_chain = (
    prompt 
    | llm
    | parser
)

company_profile = """

ABC Technologies Pvt Ltd.

We provide AI solutions,
Machine Learning applications,
Cloud services and automation.

Our expertise:

- Generative AI
- Data Analytics
- Enterprise Applications
- AI Chatbots

Experience:
5 years

"""

client_requirement = """

Client wants an AI-powered healthcare assistant.

Required features:

- Patient chatbot
- Appointment booking
- Medical FAQ system
- Analytics dashboard
- Cloud deployment

Expected timeline:
5 months

"""

proposal = proposal_chain.invoke(
    {
        "company_profile": company_profile,
        "client_requirement": client_requirement
    }
)


print(proposal)

doc = Document()

doc.add_heading(
    "Business Proposal",
    level=0
)


doc.add_paragraph(proposal)


doc.save(
    "AI_Business_Proposal.docx"
)


print("Proposal saved")

from google.colab import files

files.download(
    "AI_Business_Proposal.docx"
)
